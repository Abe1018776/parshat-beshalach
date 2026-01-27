# -*- coding: utf-8 -*-
"""
Parse beshalach document - v10
HYBRID APPROACH: Sequential within sefer + content fallback
"""

import re
import sys

def read_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def get_hebrew_words(text):
    """Extract significant Hebrew words"""
    return set(w for w in re.findall(r'[א-ת]{3,}', text) if len(w) >= 3)

def normalize_sefer(name):
    """Normalize sefer names"""
    name = name.strip()
    name = re.sub(r'הק\'|הקדוש|ז"ל|זי"ע|זצ"ל|⚜', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name

def extract_summaries_from_v2(filepath):
    """Extract summaries from the v2 HTML (already parsed correctly)"""
    html = read_file(filepath)
    summaries = []
    current_sefer = None
    
    sefer_pattern = re.compile(r'<div class="sefer-header">([^<]+)</div>')
    entry_pattern = re.compile(
        r'<div class="summary-entry">\s*'
        r'<span class="entry-id">\[(\d+)\]</span>\s*'
        r'<span class="entry-verse">([^<]+)</span>\s*'
        r'<span class="entry-summary">([^<]+)</span>\s*'
        r'<span class="entry-page">([^<]+)</span>\s*'
        r'</div>'
    )
    
    summary_section = re.search(r'<div class="source-header">מפתח ענינים</div>(.*?)<div class="divider">', html, re.DOTALL)
    if not summary_section:
        return []
    
    section_html = summary_section.group(1)
    pos = 0
    
    while pos < len(section_html):
        sefer_match = sefer_pattern.search(section_html, pos)
        entry_match = entry_pattern.search(section_html, pos)
        
        if sefer_match and (not entry_match or sefer_match.start() < entry_match.start()):
            current_sefer = sefer_match.group(1)
            pos = sefer_match.end()
        elif entry_match:
            summary_text = entry_match.group(3)
            summaries.append({
                'id': int(entry_match.group(1)),
                'sefer': normalize_sefer(current_sefer) if current_sefer else "Unknown",
                'verse': entry_match.group(2),
                'summary': summary_text,
                'page': entry_match.group(4),
                'keywords': get_hebrew_words(summary_text)
            })
            pos = entry_match.end()
        else:
            break
    
    return summaries

def extract_all_full_texts(content):
    """Extract ALL full text passages with sefer info"""
    lines = content.split('\n')
    full_texts = []
    current_sefer = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Detect sefer header with <h1>
        if '<center><h1>' in line:
            match = re.search(r'<center>\s*<h1>\s*(?:<b>)?([^<]+)(?:</b>)?\s*</h1>', line)
            if match:
                name = match.group(1).strip()
                if 'מפתח' not in name and 'ליקוטי' not in name:
                    current_sefer = normalize_sefer(name)
            i += 1
            continue
        
        if not current_sefer:
            i += 1
            continue
        
        # Detect passage: <b>verse</b>
        if line.startswith('<b>') and '</b>' in line and '<center>' not in line:
            verse_match = re.match(r'<b>([^<]+)</b>', line)
            if verse_match:
                verse = verse_match.group(1).strip()
                
                if len(verse) > 100:
                    i += 1
                    continue
                
                # Collect passage
                passage_lines = [line]
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if '<center><h1>' in next_line:
                        break
                    if next_line.startswith('<b>') and '</b>' in next_line and '<center>' not in next_line:
                        vm = re.match(r'<b>([^<]+)</b>', next_line)
                        if vm and len(vm.group(1)) < 80:
                            break
                    if next_line and '<footer>' not in next_line and '<header>' not in next_line.lower():
                        passage_lines.append(next_line)
                    j += 1
                
                text = '\n'.join(passage_lines)
                text_clean = re.sub(r'<[^>]+>', '', text)
                
                if len(text_clean) > 50:
                    full_texts.append({
                        'sefer': current_sefer,
                        'verse': verse,
                        'text': text,
                        'text_clean': text_clean,
                        'keywords': get_hebrew_words(text_clean[:800]),
                        'used': False
                    })
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return full_texts

def find_best_match(summary, full_texts):
    """Find best matching full text using verse and content matching"""
    s_verse = summary['verse']
    s_keywords = summary['keywords']
    s_summary = summary['summary']
    
    best_match = None
    best_score = 0
    best_idx = -1
    
    for idx, ft in enumerate(full_texts):
        if ft['used']:
            continue
        
        score = 0
        
        # Verse matching - primary
        ft_verse = ft['verse']
        s_verse_norm = re.sub(r'[\s\-]', '', s_verse)
        ft_verse_norm = re.sub(r'[\s\-]', '', ft_verse)
        
        if s_verse_norm == ft_verse_norm:
            score += 100
        elif s_verse_norm in ft_verse_norm:
            score += 80
        elif ft_verse_norm in s_verse_norm:
            score += 70
        elif len(s_verse_norm) > 3 and len(ft_verse_norm) > 3:
            # Check first 4 chars
            if s_verse_norm[:4] == ft_verse_norm[:4]:
                score += 50
            elif s_verse_norm[:3] == ft_verse_norm[:3]:
                score += 30
        
        # Content matching - check if summary words appear in full text
        if s_keywords:
            ft_text_start = ft['text_clean'][:600]
            matches = sum(1 for w in s_keywords if w in ft_text_start)
            content_score = (matches / len(s_keywords)) * 60
            score += content_score
        
        # Check for key phrases from summary in full text
        key_phrases = re.findall(r'[א-ת]{4,}', s_summary)[:10]
        if key_phrases:
            phrase_matches = sum(1 for p in key_phrases if p in ft['text_clean'][:500])
            score += (phrase_matches / len(key_phrases)) * 40
        
        if score > best_score:
            best_score = score
            best_match = ft
            best_idx = idx
    
    # Require minimum score
    if best_match and best_score >= 40:
        return best_match, best_idx, best_score
    
    return None, -1, 0

def match_all(summaries, full_texts):
    """Match all summaries to full texts"""
    matched = []
    
    for summary in summaries:
        ft, idx, score = find_best_match(summary, full_texts)
        
        if ft:
            full_texts[idx]['used'] = True
            matched.append({
                'summary': summary,
                'full_text': ft,
                'score': score
            })
        else:
            matched.append({
                'summary': summary,
                'full_text': None,
                'score': 0
            })
    
    return matched

def clean_html(text):
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def generate_html(matched):
    html = '''<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>פרשת בשלח - ליקוטי ספרי חסידות</title>
    <style>
        @page { size: A4; margin: 2cm; }
        body {
            font-family: 'Frank Ruehl', 'David', 'Times New Roman', serif;
            font-size: 14pt;
            line-height: 1.8;
            direction: rtl;
            text-align: justify;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #fff;
        }
        .summary-entry { display: flex; align-items: baseline; margin: 8px 0; gap: 10px; }
        .entry-id { font-weight: bold; min-width: 35px; }
        .entry-verse { font-weight: bold; color: #333; min-width: 80px; }
        .entry-summary { flex: 1; }
        .entry-page { font-weight: bold; min-width: 45px; text-align: left; }
        .divider { border-top: 2px solid #000; margin: 30px 0; }
        .source-header { text-align: center; font-weight: bold; font-size: 20pt; margin: 25px 0; }
        .sefer-header { text-align: center; font-weight: bold; font-size: 16pt; margin: 25px 0 15px 0; background: #f0f0f0; padding: 10px; border-radius: 4px; }
        .quote-entry { margin-bottom: 30px; padding: 20px; border: 1px solid #ccc; border-radius: 6px; background: #fafafa; }
        .quote-header { display: flex; gap: 20px; margin-bottom: 12px; font-weight: bold; font-size: 15pt; border-bottom: 2px solid #ddd; padding-bottom: 10px; }
        .quote-text { text-align: justify; line-height: 2; font-size: 14pt; }
        .no-match { color: #999; font-style: italic; }
        h1 { text-align: center; font-size: 24pt; margin-bottom: 5px; }
        h2 { text-align: center; font-size: 18pt; margin-top: 5px; color: #555; }
    </style>
</head>
<body>
    <h1>ליקוטי ספרי חסידות</h1>
    <h2>פרשת בשלח</h2>
    <div class="divider"></div>
    <section>
        <div class="source-header">מפתח ענינים</div>
'''
    
    current_sefer = None
    for entry in matched:
        s = entry['summary']
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        html += f'''        <div class="summary-entry">
            <span class="entry-id">[{s['id']}]</span>
            <span class="entry-verse">{s['verse']}</span>
            <span class="entry-summary">{s['summary']}</span>
            <span class="entry-page">{s['page']}</span>
        </div>\n'''
    
    html += '''    </section>
    <div class="divider"></div>
    <section>
        <div class="source-header">מקור השפע</div>
'''
    
    current_sefer = None
    for entry in matched:
        s = entry['summary']
        ft = entry['full_text']
        
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        
        if ft:
            text = clean_html(ft['text'])
            if len(text) > 3000:
                text = text[:3000] + '\n\n[...]'
        else:
            text = '<span class="no-match">[לא נמצא טקסט מתאים]</span>'
        
        html += f'''        <div class="quote-entry">
            <div class="quote-header">
                <span>[{s['id']}]</span>
                <span>{s['verse']}</span>
                <span>עמ\' {s['page']}</span>
            </div>
            <div class="quote-text">{text}</div>
        </div>\n'''
    
    html += '''    </section>
</body>
</html>'''
    return html

def generate_word(matched, output):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def set_rtl(p):
        pPr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    r = title.add_run('ליקוטי ספרי חסידות - פרשת בשלח')
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = 'David'
    r._element.rPr.rFonts.set(qn('w:cs'), 'David')
    
    doc.add_paragraph()
    count = 0
    
    for i, entry in enumerate(matched):
        s = entry['summary']
        ft = entry['full_text']
        if not ft:
            continue
        
        count += 1
        
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(h)
        hr = h.add_run(f'[{s["id"]}] {s["verse"]} (עמ\' {s["page"]})')
        hr.bold = True
        hr.font.size = Pt(14)
        hr.font.name = 'David'
        hr._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(sp)
        sr = sp.add_run(s['summary'])
        sr.font.size = Pt(13)
        sr.font.name = 'David'
        sr._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        sep = doc.add_paragraph('_' * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sh = doc.add_paragraph()
        sh.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(sh)
        shr = sh.add_run('מקור:')
        shr.bold = True
        shr.font.size = Pt(11)
        shr.font.name = 'David'
        shr._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        ft_text = clean_html(ft['text'])
        if len(ft_text) > 2500:
            ft_text = ft_text[:2500] + ' [...]'
        
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(fp)
        fr = fp.add_run(ft_text)
        fr.font.size = Pt(11)
        fr.font.name = 'David'
        fr._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        if i < len(matched) - 1:
            doc.add_page_break()
    
    doc.save(output)
    return count

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    
    print("Reading summaries from beshalach_output_v2.html...")
    summaries = extract_summaries_from_v2('beshalach_output_v2.html')
    print(f"Found {len(summaries)} summaries")
    
    print("\nReading full texts from extracted_beshalach.txt...")
    content = read_file('extracted_beshalach.txt')
    full_texts = extract_all_full_texts(content)
    print(f"Found {len(full_texts)} full text passages")
    
    print("\nMatching summaries to full texts...")
    matched = match_all(summaries, full_texts)
    
    matched_count = sum(1 for m in matched if m['full_text'])
    print(f"\nMatched: {matched_count}/{len(summaries)} ({100*matched_count/len(summaries):.1f}%)")
    
    # Show matches for verification
    print("\nSample matches (first 10):")
    for m in matched[:10]:
        s = m['summary']
        ft = m['full_text']
        if ft:
            ft_verse = ft['verse'][:25]
            print(f"  [{s['id']}] {s['verse']} -> {ft_verse}... (score: {m['score']:.0f})")
        else:
            print(f"  [{s['id']}] {s['verse']} -> NO MATCH")
    
    print("\nGenerating HTML...")
    html = generate_html(matched)
    with open('beshalach_output_final.html', 'w', encoding='utf-8') as f:
        f.write(html)
    print("Saved: beshalach_output_final.html")
    
    print("\nGenerating Word document...")
    count = generate_word(matched, 'beshalach_word_final.docx')
    print(f"Saved: beshalach_word_final.docx ({count} pages)")

if __name__ == '__main__':
    main()
