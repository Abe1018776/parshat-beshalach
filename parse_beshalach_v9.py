# -*- coding: utf-8 -*-
"""
Parse beshalach document - v9
SEQUENTIAL MATCHING: Match summaries to full texts in order within each sefer
"""

import re
import sys

def read_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def normalize_sefer_name(name):
    """Normalize sefer names for matching"""
    name = name.strip()
    # Remove common variations
    name = re.sub(r'הק\'|הקדוש|ז"ל|זי"ע|זצ"ל', '', name)
    name = re.sub(r'\s+', ' ', name).strip()
    return name

def parse_toc(content):
    """Parse table of contents to get sefer->page mapping"""
    toc = {}
    lines = content.split('\n')
    
    for line in lines:
        # TOC format: <b>sefer name</b>......page
        if '<b>' in line and '....' in line:
            match = re.match(r'<b>([^<]+)</b>[\.]+\s*([א-ת\'"]+)', line.strip())
            if match:
                sefer = normalize_sefer_name(match.group(1))
                page = match.group(2).strip()
                toc[page] = sefer
    
    return toc

def extract_summaries_grouped(content, toc):
    """Extract summaries grouped by sefer, in order"""
    lines = content.split('\n')
    summaries_by_sefer = {}
    all_summaries = []
    current_sefer = None
    
    # Find where summaries start (after TOC, before full texts)
    in_summaries = False
    entry_id = 0
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Stop at full text section
        if '<center><h1><b>' in line and 'מפתח' not in line:
            break
        
        # Detect sefer headers in summaries: <center><b>name</b></center>
        if '<center><b>' in line and '</b></center>' in line and '<h1>' not in line:
            match = re.search(r'<center><b>([^<]+)</b></center>', line)
            if match:
                name = match.group(1).strip()
                if name not in ['בעזהשי"ת', 'עם מפתח ענינים', 'פרשת בשלח']:
                    current_sefer = normalize_sefer_name(name)
                    if current_sefer not in summaries_by_sefer:
                        summaries_by_sefer[current_sefer] = []
            i += 1
            continue
        
        # Detect summary entries: <b>verse</b>-summary...page
        if line.startswith('<b>') and '</b>' in line and '-' in line:
            # Skip TOC entries
            if '....' in line and line.count('.') > 10:
                i += 1
                continue
            
            # Extract verse
            verse_match = re.match(r'<b>([^<]+)</b>', line)
            if not verse_match:
                i += 1
                continue
            
            verse_raw = verse_match.group(1).strip()
            verse = verse_raw.rstrip('-').strip()
            
            # Get text after </b>
            after_verse = line[verse_match.end():].strip()
            after_verse = after_verse.lstrip('-').strip()
            
            # Skip if too short (likely TOC remnant)
            if len(after_verse) < 20:
                i += 1
                continue
            
            # Collect full summary (may span multiple lines)
            summary_lines = [after_verse]
            page = None
            
            # Check for page at end
            page_match = re.search(r'[\.]{2,}\s*([א-ת\'"]+)\s*$', after_verse)
            if page_match:
                page = page_match.group(1)
                summary_lines[0] = re.sub(r'[\.]{2,}\s*[א-ת\'"]+\s*$', '', after_verse).strip()
            
            # Continue to next lines if no page yet
            j = i + 1
            while j < len(lines) and not page:
                next_line = lines[j].strip()
                if next_line.startswith('<b>') or next_line.startswith('<center>'):
                    break
                if next_line:
                    page_match = re.search(r'[\.]{2,}\s*([א-ת\'"]+)\s*$', next_line)
                    if page_match:
                        page = page_match.group(1)
                        summary_lines.append(re.sub(r'[\.]{2,}\s*[א-ת\'"]+\s*$', '', next_line).strip())
                        j += 1
                        break
                    else:
                        summary_lines.append(next_line)
                j += 1
            
            if page:
                entry_id += 1
                summary_text = ' '.join(summary_lines).strip()
                
                # Determine sefer from current_sefer or from page->TOC mapping
                sefer = current_sefer
                if not sefer and page in toc:
                    sefer = toc[page]
                if not sefer:
                    sefer = "Unknown"
                
                entry = {
                    'id': entry_id,
                    'sefer': sefer,
                    'verse': verse,
                    'summary': summary_text,
                    'page': page
                }
                
                all_summaries.append(entry)
                
                if sefer not in summaries_by_sefer:
                    summaries_by_sefer[sefer] = []
                summaries_by_sefer[sefer].append(entry)
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return all_summaries, summaries_by_sefer

def extract_full_texts_grouped(content):
    """Extract full texts grouped by sefer, in order"""
    lines = content.split('\n')
    full_texts_by_sefer = {}
    current_sefer = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Detect sefer headers with <h1> (full text section)
        if '<center><h1>' in line:
            match = re.search(r'<center>\s*<h1>\s*(?:<b>)?([^<]+)(?:</b>)?\s*</h1>', line)
            if match:
                name = match.group(1).strip()
                if 'מפתח' not in name and 'ליקוטי' not in name:
                    current_sefer = normalize_sefer_name(name)
                    if current_sefer not in full_texts_by_sefer:
                        full_texts_by_sefer[current_sefer] = []
            i += 1
            continue
        
        if not current_sefer:
            i += 1
            continue
        
        # Detect passage start: <b>verse</b>
        if line.startswith('<b>') and '</b>' in line and '<center>' not in line:
            verse_match = re.match(r'<b>([^<]+)</b>', line)
            if verse_match:
                verse = verse_match.group(1).strip()
                
                # Skip if too long (not a verse reference)
                if len(verse) > 100:
                    i += 1
                    continue
                
                # Collect full passage
                passage_lines = [line]
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    # Stop at next sefer header
                    if '<center><h1>' in next_line:
                        break
                    # Stop at next passage
                    if next_line.startswith('<b>') and '</b>' in next_line and '<center>' not in next_line:
                        vm = re.match(r'<b>([^<]+)</b>', next_line)
                        if vm and len(vm.group(1)) < 80:
                            break
                    if next_line and '<footer>' not in next_line and '<header>' not in next_line.lower():
                        passage_lines.append(next_line)
                    j += 1
                
                text = '\n'.join(passage_lines)
                text_clean = re.sub(r'<[^>]+>', '', text)
                
                if len(text_clean) > 50:
                    full_texts_by_sefer[current_sefer].append({
                        'sefer': current_sefer,
                        'verse': verse,
                        'text': text,
                        'text_clean': text_clean
                    })
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return full_texts_by_sefer

def find_sefer_match(summary_sefer, full_text_sefers):
    """Find matching sefer name accounting for variations"""
    if summary_sefer in full_text_sefers:
        return summary_sefer
    
    # Try partial matching
    s_norm = normalize_sefer_name(summary_sefer).lower()
    for ft_sefer in full_text_sefers:
        ft_norm = normalize_sefer_name(ft_sefer).lower()
        if s_norm in ft_norm or ft_norm in s_norm:
            return ft_sefer
        # Check first word match
        s_first = s_norm.split()[0] if s_norm.split() else ""
        ft_first = ft_norm.split()[0] if ft_norm.split() else ""
        if len(s_first) > 3 and s_first == ft_first:
            return ft_sefer
    
    return None

def match_sequential(all_summaries, summaries_by_sefer, full_texts_by_sefer):
    """Match summaries to full texts sequentially within each sefer"""
    matched = []
    sefer_indices = {}  # Track current index per sefer
    
    for summary in all_summaries:
        sefer = summary['sefer']
        
        # Find matching sefer in full texts
        ft_sefer = find_sefer_match(sefer, full_texts_by_sefer.keys())
        
        if ft_sefer and ft_sefer in full_texts_by_sefer:
            # Get current index for this sefer
            if ft_sefer not in sefer_indices:
                sefer_indices[ft_sefer] = 0
            
            idx = sefer_indices[ft_sefer]
            ft_list = full_texts_by_sefer[ft_sefer]
            
            if idx < len(ft_list):
                matched.append({
                    'summary': summary,
                    'full_text': ft_list[idx],
                    'match_type': 'sequential'
                })
                sefer_indices[ft_sefer] = idx + 1
            else:
                matched.append({
                    'summary': summary,
                    'full_text': None,
                    'match_type': 'no_more_texts'
                })
        else:
            matched.append({
                'summary': summary,
                'full_text': None,
                'match_type': 'no_sefer_match'
            })
    
    return matched

def clean_html(text):
    """Clean HTML tags"""
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def generate_html(matched_entries):
    """Generate HTML output"""
    html = '''<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>פרשת בשלח - ליקוטי ספרי חסידות</title>
    <style>
        @page { size: A4; margin: 2cm; }
        body {
            font-family: 'Frank Ruehl', 'David', 'Times New Roman', serif;
            font-size: 14pt;
            line-height: 1.8;
            direction: rtl;
            text-align: justify;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #fff;
        }
        .summary-entry {
            display: flex;
            align-items: baseline;
            margin: 8px 0;
            gap: 10px;
        }
        .entry-id { font-weight: bold; min-width: 35px; }
        .entry-verse { font-weight: bold; color: #333; min-width: 80px; }
        .entry-summary { flex: 1; }
        .entry-page { font-weight: bold; min-width: 45px; text-align: left; }
        .divider { border-top: 2px solid #000; margin: 30px 0; }
        .source-header {
            text-align: center;
            font-weight: bold;
            font-size: 20pt;
            margin: 25px 0;
            color: #222;
        }
        .sefer-header {
            text-align: center;
            font-weight: bold;
            font-size: 16pt;
            margin: 25px 0 15px 0;
            background: #f0f0f0;
            padding: 10px;
            border-radius: 4px;
        }
        .quote-entry {
            margin-bottom: 30px;
            padding: 20px;
            border: 1px solid #ccc;
            border-radius: 6px;
            background: #fafafa;
        }
        .quote-header {
            display: flex;
            gap: 20px;
            margin-bottom: 12px;
            font-weight: bold;
            font-size: 15pt;
            border-bottom: 2px solid #ddd;
            padding-bottom: 10px;
            color: #444;
        }
        .quote-text { 
            text-align: justify; 
            line-height: 2;
            font-size: 14pt;
        }
        .no-match {
            color: #999;
            font-style: italic;
        }
        h1 { text-align: center; font-size: 24pt; margin-bottom: 5px; }
        h2 { text-align: center; font-size: 18pt; margin-top: 5px; color: #555; }
    </style>
</head>
<body>
    <h1>ליקוטי ספרי חסידות</h1>
    <h2>פרשת בשלח</h2>
    <div class="divider"></div>
    
    <section>
        <div class="source-header">מפתח ענינים</div>
'''
    
    # Summaries section
    current_sefer = None
    for entry in matched_entries:
        s = entry['summary']
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        html += f'''        <div class="summary-entry">
            <span class="entry-id">[{s['id']}]</span>
            <span class="entry-verse">{s['verse']}</span>
            <span class="entry-summary">{s['summary']}</span>
            <span class="entry-page">{s['page']}</span>
        </div>\n'''
    
    html += '''    </section>
    <div class="divider"></div>
    
    <section>
        <div class="source-header">מקור השפע</div>
'''
    
    # Full texts section
    current_sefer = None
    for entry in matched_entries:
        s = entry['summary']
        ft = entry['full_text']
        
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        
        if ft:
            text = clean_html(ft['text'])
            if len(text) > 3000:
                text = text[:3000] + '\n\n[...]'
        else:
            text = f'<span class="no-match">[לא נמצא טקסט מתאים - {entry["match_type"]}]</span>'
        
        html += f'''        <div class="quote-entry">
            <div class="quote-header">
                <span>[{s['id']}]</span>
                <span>{s['verse']}</span>
                <span>עמ\' {s['page']}</span>
            </div>
            <div class="quote-text">{text}</div>
        </div>\n'''
    
    html += '''    </section>
</body>
</html>'''
    
    return html

def generate_word_doc(matched_entries, output_path):
    """Generate Word document"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    run = title.add_run('ליקוטי ספרי חסידות - פרשת בשלח')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'David'
    run._element.rPr.rFonts.set(qn('w:cs'), 'David')
    
    doc.add_paragraph()
    
    count = 0
    for i, entry in enumerate(matched_entries):
        s = entry['summary']
        ft = entry['full_text']
        
        if not ft:
            continue
        
        count += 1
        
        # Summary header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(header_para)
        
        header_run = header_para.add_run(f'[{s["id"]}] {s["verse"]} ')
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = 'David'
        header_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        page_run = header_para.add_run(f'(עמ\' {s["page"]})')
        page_run.font.size = Pt(12)
        page_run.font.name = 'David'
        page_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Summary text
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(summary_para)
        
        summary_run = summary_para.add_run(s['summary'])
        summary_run.font.size = Pt(13)
        summary_run.font.name = 'David'
        summary_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Separator
        sep = doc.add_paragraph('_' * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Source header
        src_header = doc.add_paragraph()
        src_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(src_header)
        src_run = src_header.add_run('מקור:')
        src_run.bold = True
        src_run.font.size = Pt(11)
        src_run.font.name = 'David'
        src_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Full text
        full_text = clean_html(ft['text'])
        if len(full_text) > 2500:
            full_text = full_text[:2500] + ' [...]'
        
        ft_para = doc.add_paragraph()
        ft_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(ft_para)
        
        ft_run = ft_para.add_run(full_text)
        ft_run.font.size = Pt(11)
        ft_run.font.name = 'David'
        ft_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Page break
        if i < len(matched_entries) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    return count

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    
    print("Reading extracted_beshalach.txt...")
    content = read_file('extracted_beshalach.txt')
    
    print("\nParsing table of contents...")
    toc = parse_toc(content)
    print(f"Found {len(toc)} seforim in TOC")
    
    print("\nExtracting summaries (grouped by sefer)...")
    all_summaries, summaries_by_sefer = extract_summaries_grouped(content, toc)
    print(f"Found {len(all_summaries)} summaries in {len(summaries_by_sefer)} seforim:")
    for sefer, items in summaries_by_sefer.items():
        print(f"  {sefer}: {len(items)} entries")
    
    print("\nExtracting full texts (grouped by sefer)...")
    full_texts_by_sefer = extract_full_texts_grouped(content)
    print(f"Found full texts in {len(full_texts_by_sefer)} seforim:")
    for sefer, items in full_texts_by_sefer.items():
        print(f"  {sefer}: {len(items)} passages")
    
    print("\nMatching sequentially within each sefer...")
    matched = match_sequential(all_summaries, summaries_by_sefer, full_texts_by_sefer)
    
    matched_count = sum(1 for m in matched if m['full_text'] is not None)
    print(f"\nMatched: {matched_count}/{len(all_summaries)} ({100*matched_count/len(all_summaries):.1f}%)")
    
    # Show some matches
    print("\nSample matches:")
    for m in matched[:10]:
        s = m['summary']
        ft = m['full_text']
        print(f"  [{s['id']}] {s['sefer']}: {s['verse']}")
        if ft:
            print(f"       -> {ft['verse'][:30]}... ({m['match_type']})")
        else:
            print(f"       -> NO MATCH ({m['match_type']})")
    
    # Generate outputs
    print("\nGenerating HTML...")
    html = generate_html(matched)
    with open('beshalach_output_v5.html', 'w', encoding='utf-8') as f:
        f.write(html)
    print("Saved: beshalach_output_v5.html")
    
    print("\nGenerating Word document...")
    count = generate_word_doc(matched, 'beshalach_word_output_v3.docx')
    print(f"Saved: beshalach_word_output_v3.docx ({count} pages)")

if __name__ == '__main__':
    main()
