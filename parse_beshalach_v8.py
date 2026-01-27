# -*- coding: utf-8 -*-
"""
Parse beshalach document - v8
Fixed matching: Use CONTENT-BASED matching as primary criteria
"""

import re
import sys

def read_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def normalize(text):
    """Normalize for comparison - remove punctuation and spaces"""
    return re.sub(r'[\s"\'\-־–—׳״\.\,\:\;\!\?\(\)\[\]]', '', text).strip()

def get_hebrew_words(text):
    """Extract Hebrew words from text"""
    return set(re.findall(r'[א-ת]{2,}', text))

def extract_summaries_from_v2_html(html):
    """Extract already-parsed summaries from v2 HTML"""
    summaries = []
    current_sefer = None
    
    sefer_pattern = re.compile(r'<div class="sefer-header">([^<]+)</div>')
    entry_pattern = re.compile(
        r'<div class="summary-entry">\s*'
        r'<span class="entry-id">\[(\d+)\]</span>\s*'
        r'<span class="entry-verse">([^<]+)</span>\s*'
        r'<span class="entry-summary">([^<]+)</span>\s*'
        r'<span class="entry-page">([^<]+)</span>\s*'
        r'</div>'
    )
    
    # Find summary section
    summary_section = re.search(r'<div class="source-header">מפתח ענינים</div>(.*?)<div class="divider">', html, re.DOTALL)
    if not summary_section:
        return []
    
    section_html = summary_section.group(1)
    pos = 0
    
    while pos < len(section_html):
        sefer_match = sefer_pattern.search(section_html, pos)
        entry_match = entry_pattern.search(section_html, pos)
        
        if sefer_match and (not entry_match or sefer_match.start() < entry_match.start()):
            current_sefer = sefer_match.group(1)
            pos = sefer_match.end()
        elif entry_match:
            summary_text = entry_match.group(3)
            summaries.append({
                'id': int(entry_match.group(1)),
                'sefer': current_sefer,
                'verse': entry_match.group(2),
                'summary': summary_text,
                'page': entry_match.group(4),
                'keywords': get_hebrew_words(summary_text)
            })
            pos = entry_match.end()
        else:
            break
    
    return summaries

def extract_full_texts(content):
    """Extract full text passages from מקור השפע section"""
    lines = content.split('\n')
    full_texts = []
    current_sefer = None
    in_full_text = False
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Detect sefer headers with <h1> (full text section)
        if '<center><h1>' in line_stripped:
            match = re.search(r'<center>\s*<h1>\s*(?:<b>)?([^<]+)(?:</b>)?\s*</h1>', line_stripped)
            if match:
                name = match.group(1).strip()
                if 'מפתח' not in name and 'ליקוטי' not in name:
                    current_sefer = name
                    in_full_text = True
                    continue
        
        if not in_full_text or not current_sefer:
            continue
        
        # Detect passage start: <b>verse text</b>
        if line_stripped.startswith('<b>') and '</b>' in line_stripped:
            verse_match = re.match(r'<b>([^<]+)</b>', line_stripped)
            if verse_match and '<center>' not in line_stripped:
                verse = verse_match.group(1).strip()
                
                # Skip if too long (likely not a verse reference)
                if len(verse) > 100:
                    continue
                
                # Collect full passage
                passage_lines = [line_stripped]
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    # Stop at next sefer or next passage
                    if '<center><h1>' in next_line:
                        break
                    if next_line.startswith('<b>') and '</b>' in next_line and '<center>' not in next_line:
                        # Check if this is a new passage (short verse reference)
                        vm = re.match(r'<b>([^<]+)</b>', next_line)
                        if vm and len(vm.group(1)) < 80:
                            break
                    if next_line and '<footer>' not in next_line and '<header>' not in next_line.lower():
                        passage_lines.append(next_line)
                    j += 1
                
                text = '\n'.join(passage_lines)
                # Clean HTML
                text_clean = re.sub(r'<[^>]+>', '', text)
                
                if len(text_clean) > 100:
                    full_texts.append({
                        'sefer': current_sefer,
                        'verse': verse,
                        'text': text,
                        'text_clean': text_clean,
                        'keywords': get_hebrew_words(text_clean[:1000])
                    })
    
    return full_texts

def content_match_score(summary, full_text):
    """Calculate content match score based on keyword overlap"""
    summary_keywords = summary['keywords']
    ft_keywords = full_text['keywords']
    
    if not summary_keywords:
        return 0
    
    # Count how many summary keywords appear in full text
    matches = len(summary_keywords & ft_keywords)
    score = matches / len(summary_keywords)
    
    # Also check first 500 chars of full text for summary words
    ft_start = normalize(full_text['text_clean'][:500])
    summary_norm = normalize(summary['summary'])
    
    # Check for significant phrase matches
    summary_words = [w for w in re.findall(r'[א-ת]{3,}', summary['summary']) if len(w) >= 3]
    phrase_matches = sum(1 for w in summary_words if w in ft_start)
    if summary_words:
        phrase_score = phrase_matches / len(summary_words)
        score = max(score, phrase_score)
    
    return score

def match_summaries_to_full_texts(summaries, full_texts):
    """Match each summary to its best full text using content matching"""
    matched = []
    used_indices = set()
    
    for summary in summaries:
        best_match = None
        best_score = 0
        best_idx = -1
        
        for idx, ft in enumerate(full_texts):
            if idx in used_indices:
                continue
            
            # Calculate content match score
            content_score = content_match_score(summary, ft)
            
            # Bonus for verse match
            verse_bonus = 0
            s_verse = normalize(summary['verse'])
            ft_verse = normalize(ft['verse'])
            if s_verse and ft_verse:
                if s_verse in ft_verse or ft_verse in s_verse:
                    verse_bonus = 0.2
                elif s_verse[:3] == ft_verse[:3]:
                    verse_bonus = 0.1
            
            total_score = content_score + verse_bonus
            
            if total_score > best_score:
                best_score = total_score
                best_match = ft
                best_idx = idx
        
        if best_match and best_score >= 0.25:
            used_indices.add(best_idx)
            matched.append({
                'summary': summary,
                'full_text': best_match,
                'score': best_score
            })
        else:
            matched.append({
                'summary': summary,
                'full_text': None,
                'score': 0
            })
    
    return matched

def clean_html(text):
    """Clean HTML tags for display"""
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def generate_html(matched_entries):
    """Generate HTML output"""
    html = '''<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
    <meta charset="UTF-8">
    <title>פרשת בשלח - ליקוטי ספרי חסידות</title>
    <style>
        @page { size: A4; margin: 2cm; }
        body {
            font-family: 'Frank Ruehl', 'David', 'Times New Roman', serif;
            font-size: 14pt;
            line-height: 1.8;
            direction: rtl;
            text-align: justify;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: #fff;
        }
        .summary-entry {
            display: flex;
            align-items: baseline;
            margin: 8px 0;
            gap: 10px;
        }
        .entry-id { font-weight: bold; min-width: 35px; }
        .entry-verse { font-weight: bold; color: #333; min-width: 80px; }
        .entry-summary { flex: 1; }
        .entry-page { font-weight: bold; min-width: 45px; text-align: left; }
        .divider { border-top: 2px solid #000; margin: 30px 0; }
        .source-header {
            text-align: center;
            font-weight: bold;
            font-size: 20pt;
            margin: 25px 0;
            color: #222;
        }
        .sefer-header {
            text-align: center;
            font-weight: bold;
            font-size: 16pt;
            margin: 25px 0 15px 0;
            background: #f0f0f0;
            padding: 10px;
            border-radius: 4px;
        }
        .quote-entry {
            margin-bottom: 30px;
            padding: 20px;
            border: 1px solid #ccc;
            border-radius: 6px;
            background: #fafafa;
        }
        .quote-header {
            display: flex;
            gap: 20px;
            margin-bottom: 12px;
            font-weight: bold;
            font-size: 15pt;
            border-bottom: 2px solid #ddd;
            padding-bottom: 10px;
            color: #444;
        }
        .quote-text { 
            text-align: justify; 
            line-height: 2;
            font-size: 14pt;
        }
        .match-score {
            color: #666;
            font-size: 10pt;
        }
        .no-match {
            color: #999;
            font-style: italic;
        }
        h1 { text-align: center; font-size: 24pt; margin-bottom: 5px; }
        h2 { text-align: center; font-size: 18pt; margin-top: 5px; color: #555; }
    </style>
</head>
<body>
    <h1>ליקוטי ספרי חסידות</h1>
    <h2>פרשת בשלח</h2>
    <div class="divider"></div>
    
    <section>
        <div class="source-header">מפתח ענינים</div>
'''
    
    # Summaries section
    current_sefer = None
    for entry in matched_entries:
        s = entry['summary']
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        html += f'''        <div class="summary-entry">
            <span class="entry-id">[{s['id']}]</span>
            <span class="entry-verse">{s['verse']}</span>
            <span class="entry-summary">{s['summary']}</span>
            <span class="entry-page">{s['page']}</span>
        </div>\n'''
    
    html += '''    </section>
    <div class="divider"></div>
    
    <section>
        <div class="source-header">מקור השפע</div>
'''
    
    # Full texts section
    current_sefer = None
    for entry in matched_entries:
        s = entry['summary']
        ft = entry['full_text']
        
        if s['sefer'] != current_sefer:
            current_sefer = s['sefer']
            html += f'        <div class="sefer-header">{current_sefer}</div>\n'
        
        if ft:
            text = clean_html(ft['text'])
            if len(text) > 3000:
                text = text[:3000] + '\n\n[...]'
            score_pct = int(entry['score'] * 100)
        else:
            text = f'<span class="no-match">[לא נמצא טקסט מתאים]</span>'
            score_pct = 0
        
        html += f'''        <div class="quote-entry">
            <div class="quote-header">
                <span>[{s['id']}]</span>
                <span>{s['verse']}</span>
                <span>עמ\' {s['page']}</span>
            </div>
            <div class="quote-text">{text}</div>
        </div>\n'''
    
    html += '''    </section>
</body>
</html>'''
    
    return html

def generate_word_doc(matched_entries, output_path):
    """Generate Word document with summaries and full texts"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    run = title.add_run('ליקוטי ספרי חסידות - פרשת בשלח')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'David'
    run._element.rPr.rFonts.set(qn('w:cs'), 'David')
    
    doc.add_paragraph()
    
    count = 0
    for i, entry in enumerate(matched_entries):
        s = entry['summary']
        ft = entry['full_text']
        
        if not ft:
            continue
        
        count += 1
        
        # Summary header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(header_para)
        
        header_run = header_para.add_run(f'[{s["id"]}] {s["verse"]} ')
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = 'David'
        header_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        page_run = header_para.add_run(f'(עמ\' {s["page"]})')
        page_run.font.size = Pt(12)
        page_run.font.name = 'David'
        page_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Summary text
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(summary_para)
        
        summary_run = summary_para.add_run(s['summary'])
        summary_run.font.size = Pt(13)
        summary_run.font.name = 'David'
        summary_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Separator
        sep = doc.add_paragraph('_' * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Source header
        src_header = doc.add_paragraph()
        src_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(src_header)
        src_run = src_header.add_run('מקור:')
        src_run.bold = True
        src_run.font.size = Pt(11)
        src_run.font.name = 'David'
        src_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Full text
        full_text = clean_html(ft['text'])
        if len(full_text) > 2500:
            full_text = full_text[:2500] + ' [...]'
        
        ft_para = doc.add_paragraph()
        ft_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(ft_para)
        
        ft_run = ft_para.add_run(full_text)
        ft_run.font.size = Pt(11)
        ft_run.font.name = 'David'
        ft_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Page break
        if i < len(matched_entries) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    return count

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    
    print("Reading beshalach_output_v2.html for summaries...")
    v2_html = read_file('beshalach_output_v2.html')
    
    print("Extracting summaries from v2 HTML...")
    summaries = extract_summaries_from_v2_html(v2_html)
    
    print("Reading extracted_beshalach.txt for full texts...")
    content = read_file('extracted_beshalach.txt')
    print(f"Found {len(summaries)} summaries")
    
    print("\nExtracting full texts...")
    full_texts = extract_full_texts(content)
    print(f"Found {len(full_texts)} full text passages")
    
    print("\nMatching summaries to full texts (content-based)...")
    matched = match_summaries_to_full_texts(summaries, full_texts)
    
    matched_count = sum(1 for m in matched if m['full_text'] is not None)
    print(f"Matched: {matched_count}/{len(summaries)} ({100*matched_count/len(summaries):.1f}%)")
    
    # Show some matches for verification
    print("\nSample matches:")
    for m in matched[:5]:
        s = m['summary']
        ft = m['full_text']
        print(f"  [{s['id']}] {s['verse']}: {s['summary'][:40]}...")
        if ft:
            print(f"       -> {ft['text_clean'][:60]}... (score: {m['score']:.2f})")
        else:
            print(f"       -> NO MATCH")
    
    # Generate HTML
    print("\nGenerating HTML...")
    html = generate_html(matched)
    with open('beshalach_output_v4.html', 'w', encoding='utf-8') as f:
        f.write(html)
    print("Saved: beshalach_output_v4.html")
    
    # Generate Word doc
    print("\nGenerating Word document...")
    count = generate_word_doc(matched, 'beshalach_word_output_v2.docx')
    print(f"Saved: beshalach_word_output_v2.docx ({count} pages)")

if __name__ == '__main__':
    main()
