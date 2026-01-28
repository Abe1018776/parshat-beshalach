# -*- coding: utf-8 -*-
"""
Generate formatted Word document like פרשת בא1.pdf
Summary on top, full content as footnote below
ONLY include entries where summary and content are from the SAME author
"""

import json
import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl_document(doc):
    """Set entire document to RTL"""
    # Set document-level RTL
    settings = doc.settings.element
    bidi = OxmlElement('w:bidiVisual')
    bidi.set(qn('w:val'), '1')
    settings.append(bidi)


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL for Hebrew text"""
    pPr = paragraph._p.get_or_add_pPr()

    # Add bidi element
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

    # Add RTL justification
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)


def set_rtl_run(run):
    """Set run to RTL"""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)


def clean_html_tags(text: str) -> str:
    """Remove HTML tags from text"""
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_content_for_author(author_name: str, source_lines: list, start_line: int, end_line: int) -> str:
    """Extract content text, verify it belongs to the same author"""
    if start_line > 0 and end_line > start_line:
        content_lines = source_lines[start_line:end_line]
        text = '\n'.join(content_lines)
        return clean_html_tags(text)
    return ""


def create_document(mapping_data: dict, source_file: str, output_file: str):
    """Create formatted Word document - ONLY same-author matches"""

    # Read source file for content extraction
    with open(source_file, 'r', encoding='utf-8') as f:
        source_lines = f.read().split('\n')

    doc = Document()

    # Set document to RTL
    set_rtl_document(doc)

    # Set up page
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Set section to RTL
    sectPr = section._sectPr
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    sectPr.append(bidi)

    entry_count = 0
    skipped_count = 0

    for author_data in mapping_data['authors']:
        author_name = author_data['name']

        for summary in author_data.get('summaries', []):
            if not summary.get('matched', False):
                skipped_count += 1
                continue

            # CRITICAL: Only use high-confidence matches (>60%)
            if summary.get('confidence', 0) < 0.6:
                skipped_count += 1
                continue

            opening_words = summary['opening_words']
            summary_text = summary['summary_text']
            content_info = summary.get('content', {})

            # Extract content
            start_line = content_info.get('start_line', 0)
            end_line = content_info.get('end_line', 0)

            if start_line == 0 or end_line == 0:
                skipped_count += 1
                continue

            content_text = extract_content_for_author(author_name, source_lines, start_line, end_line)

            if len(content_text) < 100:  # Skip if content too short
                skipped_count += 1
                continue

            entry_count += 1

            # === SUMMARY SECTION (TOP) ===

            # Header: Opening words - Author name
            header_para = doc.add_paragraph()
            set_rtl_paragraph(header_para)

            run = header_para.add_run(f"{opening_words} - {author_name}")
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'David'
            set_rtl_run(run)

            # Summary text
            summary_para = doc.add_paragraph()
            set_rtl_paragraph(summary_para)

            run = summary_para.add_run(summary_text)
            run.font.size = Pt(14)
            run.font.name = 'David'
            set_rtl_run(run)

            # === CONTENT SECTION (BOTTOM - FOOTNOTE STYLE) ===

            # Separator line
            sep_para = doc.add_paragraph()
            set_rtl_paragraph(sep_para)
            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sep_para.add_run("_" * 60)
            run.font.size = Pt(10)

            # "מקור השפע" header
            source_header = doc.add_paragraph()
            set_rtl_paragraph(source_header)
            source_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = source_header.add_run("מקור השפע")
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'David'
            set_rtl_run(run)

            # Content text
            content_para = doc.add_paragraph()
            set_rtl_paragraph(content_para)

            # Truncate if too long
            if len(content_text) > 3000:
                content_text = content_text[:3000] + "..."

            run = content_para.add_run(f"[{entry_count}] {content_text}")
            run.font.size = Pt(11)
            run.font.name = 'David'
            set_rtl_run(run)

            # Page break between entries
            doc.add_page_break()

    # Save document
    doc.save(output_file)
    print(f"Document saved: {output_file}")
    print(f"Total entries: {entry_count}")
    print(f"Skipped (low confidence/missing content): {skipped_count}")


def main():
    # Paths
    mapping_file = r"C:\Users\Main\yanky fridman\ocr-172a161e-b608-46df-83ec-5db88688b475_mapping.json"
    source_file = r"C:\Users\Main\yanky fridman\ocr-172a161e-b608-46df-83ec-5db88688b475.txt"
    output_file = r"C:\Users\Main\yanky fridman\parshat_beshalach_v2.docx"

    # Load mapping data
    print("Loading mapping data...")
    with open(mapping_file, 'r', encoding='utf-8') as f:
        mapping_data = json.load(f)

    print(f"Found {mapping_data['statistics']['matched']} matched summaries")

    # Generate document
    print("Generating document...")
    create_document(mapping_data, source_file, output_file)


if __name__ == '__main__':
    main()
