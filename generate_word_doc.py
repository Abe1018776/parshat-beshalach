# -*- coding: utf-8 -*-
"""
Generate Word document from beshalach_output_v3.html
Each page: Summary on top, Full text as footnote at bottom
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def read_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def extract_entries_from_html(html):
    """Extract matched summary + full text pairs from v3 HTML"""
    entries = []
    
    # Extract summaries
    summary_pattern = re.compile(
        r'<div class="summary-entry">\s*'
        r'<span class="entry-id">\[(\d+)\]</span>\s*'
        r'<span class="entry-verse">([^<]+)</span>\s*'
        r'<span class="entry-summary">([^<]+)</span>\s*'
        r'<span class="entry-page">([^<]+)</span>\s*'
        r'</div>',
        re.DOTALL
    )
    
    # Get summaries
    summaries = {}
    for match in summary_pattern.finditer(html):
        entry_id = match.group(1)
        summaries[entry_id] = {
            'id': entry_id,
            'verse': match.group(2).strip(),
            'summary': match.group(3).strip(),
            'page': match.group(4).strip()
        }
    
    # Extract full texts using simpler approach - find quote-entry blocks
    quote_blocks = re.findall(
        r'<div class="quote-entry">(.*?)</div>\s*</div>',
        html, re.DOTALL
    )
    
    for block in quote_blocks:
        # Extract ID from header
        id_match = re.search(r'<span>\[(\d+)\]</span>', block)
        if not id_match:
            continue
        entry_id = id_match.group(1)
        
        if entry_id not in summaries:
            continue
        
        # Extract text from quote-text div
        text_match = re.search(r'<div class="quote-text">(.*)', block, re.DOTALL)
        if not text_match:
            continue
        
        full_text = text_match.group(1).strip()
        # Clean HTML tags
        full_text = re.sub(r'<div class="no-match">', '', full_text)
        full_text = re.sub(r'</div>', '', full_text)
        full_text = re.sub(r'<[^>]+>', '', full_text)
        full_text = full_text.strip()
        
        # Only include if we have actual full text (not just summary repeated)
        if len(full_text) > len(summaries[entry_id]['summary']) + 50:
            entries.append({
                'id': entry_id,
                'verse': summaries[entry_id]['verse'],
                'summary': summaries[entry_id]['summary'],
                'page': summaries[entry_id]['page'],
                'full_text': full_text
            })
    
    return entries

def set_rtl(paragraph):
    """Set paragraph to RTL"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def add_footnote(paragraph, footnote_text, doc):
    """Add footnote to paragraph"""
    # Get or create footnotes part
    footnotes_part = doc.part.footnotes_part
    if footnotes_part is None:
        # Create footnotes if not exist
        from docx.oxml.ns import nsmap
        footnotes_xml = (
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            '</w:footnotes>'
        )
    
    # Create footnote reference in main text
    run = paragraph.add_run()
    footnote_ref = OxmlElement('w:footnoteReference')
    footnote_id = str(len(doc.element.body) + 1)
    footnote_ref.set(qn('w:id'), footnote_id)
    run._r.append(footnote_ref)
    
    return footnote_id

def create_word_document(entries, output_path):
    """Create Word document with summaries and footnotes"""
    doc = Document()
    
    # Set document to RTL
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl(title)
    run = title.add_run('ליקוטי ספרי חסידות - פרשת בשלח')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'David'
    run._element.rPr.rFonts.set(qn('w:cs'), 'David')
    
    doc.add_paragraph()
    
    for i, entry in enumerate(entries):
        # Summary section (top of page area for this entry)
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(summary_para)
        
        # Entry header
        header_run = summary_para.add_run(f'[{entry["id"]}] {entry["verse"]} ')
        header_run.bold = True
        header_run.font.size = Pt(14)
        header_run.font.name = 'David'
        header_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Page reference
        page_run = summary_para.add_run(f'(עמ\' {entry["page"]})')
        page_run.font.size = Pt(12)
        page_run.font.name = 'David'
        page_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Summary text
        summary_text_para = doc.add_paragraph()
        summary_text_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(summary_text_para)
        
        summary_run = summary_text_para.add_run(entry['summary'])
        summary_run.font.size = Pt(13)
        summary_run.font.name = 'David'
        summary_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Add footnote marker
        fn_marker = summary_text_para.add_run(' *')
        fn_marker.font.superscript = True
        fn_marker.font.size = Pt(10)
        
        # Separator line
        sep = doc.add_paragraph('_' * 60)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Full text as "footnote" section at bottom
        fn_header = doc.add_paragraph()
        fn_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(fn_header)
        fn_header_run = fn_header.add_run('* מקור:')
        fn_header_run.bold = True
        fn_header_run.font.size = Pt(11)
        fn_header_run.font.name = 'David'
        fn_header_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Full text content
        full_text_para = doc.add_paragraph()
        full_text_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_rtl(full_text_para)
        
        # Limit text length for readability
        full_text = entry['full_text']
        if len(full_text) > 2500:
            full_text = full_text[:2500] + ' [...]'
        
        full_run = full_text_para.add_run(full_text)
        full_run.font.size = Pt(11)
        full_run.font.name = 'David'
        full_run._element.rPr.rFonts.set(qn('w:cs'), 'David')
        
        # Page break after each entry (except last)
        if i < len(entries) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    return len(entries)

def main():
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    
    print("Reading beshalach_output_v3.html...")
    html = read_file('beshalach_output_v3.html')
    
    print("Extracting matched entries...")
    entries = extract_entries_from_html(html)
    print(f"Found {len(entries)} entries with full text matches")
    
    output_path = 'beshalach_word_output.docx'
    print(f"\nGenerating Word document: {output_path}")
    count = create_word_document(entries, output_path)
    
    print(f"\nDone! Created {output_path} with {count} pages (one entry per page)")
    print("Each page has:")
    print("  - TOP: Summary with verse reference")
    print("  - BOTTOM: Full source text")

if __name__ == '__main__':
    main()
